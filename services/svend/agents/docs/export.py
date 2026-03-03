"""
Document Export

Export workflow results to various formats:
- Markdown
- LaTeX
- HTML
- PDF (via LaTeX or weasyprint)
"""

import json
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Optional
from datetime import datetime


@dataclass
class Section:
    """A document section."""
    title: str
    content: str
    level: int = 1
    subsections: list["Section"] = None

    def __post_init__(self):
        if self.subsections is None:
            self.subsections = []


class DocumentExporter:
    """
    Export content to various document formats.
    """

    def __init__(self, output_dir: Path = None):
        self.output_dir = output_dir or Path(".")

    def export(self, sections: list[dict], output_path: Path,
               format: str = "markdown", **kwargs) -> Path:
        """
        Export sections to a document.

        Args:
            sections: List of {"title": str, "content": str}
            output_path: Where to save
            format: markdown, latex, html, pdf
            **kwargs: title, author, date, etc.
        """
        output_path = Path(output_path)

        if format == "markdown":
            return self._export_markdown(sections, output_path, **kwargs)
        elif format == "latex":
            return self._export_latex(sections, output_path, **kwargs)
        elif format == "html":
            return self._export_html(sections, output_path, **kwargs)
        elif format == "pdf":
            return self._export_pdf(sections, output_path, **kwargs)
        elif format == "json":
            return self._export_json(sections, output_path, **kwargs)
        elif format == "docx":
            return self._export_docx(sections, output_path, **kwargs)
        else:
            raise ValueError(f"Unknown format: {format}")

    def _export_markdown(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to Markdown."""
        lines = []

        # Title
        if kwargs.get('title'):
            lines.append(f"# {kwargs['title']}")
            lines.append("")

        # Metadata
        if kwargs.get('author') or kwargs.get('date'):
            if kwargs.get('author'):
                lines.append(f"**Author:** {kwargs['author']}")
            lines.append(f"**Date:** {kwargs.get('date', datetime.now().strftime('%Y-%m-%d'))}")
            lines.append("")
            lines.append("---")
            lines.append("")

        # Sections
        for section in sections:
            title = section.get('title', 'Section')
            content = section.get('content', '')
            level = section.get('level', 2)

            lines.append(f"{'#' * level} {title}")
            lines.append("")
            lines.append(content)
            lines.append("")

        output_path.write_text("\n".join(lines))
        return output_path

    def _export_latex(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to LaTeX."""
        from .latex import LaTeXFormatter

        formatter = LaTeXFormatter()

        content_parts = []
        for section in sections:
            title = section.get('title', 'Section')
            content = section.get('content', '')
            level = section.get('level', 1)

            # Convert markdown-style content to LaTeX
            latex_content = self._markdown_to_latex(content)
            content_parts.append(formatter.section(title, latex_content, level))

        full_content = "\n\n".join(content_parts)

        document = formatter.document(
            title=kwargs.get('title', 'Document'),
            author=kwargs.get('author', ''),
            content=full_content,
            packages=kwargs.get('packages', []),
        )

        output_path.write_text(document)
        return output_path

    def _export_html(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to HTML."""
        title = kwargs.get('title', 'Document')

        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            f"<title>{title}</title>",
            "<meta charset='utf-8'>",
            "<style>",
            self._get_html_styles(),
            "</style>",
            # KaTeX for math rendering
            '<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.4/dist/katex.min.css">',
            '<script src="https://cdn.jsdelivr.net/npm/katex@0.16.4/dist/katex.min.js"></script>',
            '<script src="https://cdn.jsdelivr.net/npm/katex@0.16.4/dist/contrib/auto-render.min.js"></script>',
            "</head>",
            "<body>",
            "<article>",
        ]

        # Header
        if title:
            html_parts.append(f"<h1>{title}</h1>")

        if kwargs.get('author'):
            html_parts.append(f"<p class='author'>By {kwargs['author']}</p>")

        html_parts.append(f"<p class='date'>{kwargs.get('date', datetime.now().strftime('%Y-%m-%d'))}</p>")

        # Sections
        for section in sections:
            title = section.get('title', '')
            content = section.get('content', '')
            level = section.get('level', 2)

            html_parts.append(f"<h{level}>{title}</h{level}>")
            html_parts.append(f"<div class='section-content'>{self._markdown_to_html(content)}</div>")

        html_parts.extend([
            "</article>",
            "<script>",
            "document.addEventListener('DOMContentLoaded', function() {",
            "  renderMathInElement(document.body, {delimiters: [",
            "    {left: '$$', right: '$$', display: true},",
            "    {left: '$', right: '$', display: false},",
            "    {left: '\\\\[', right: '\\\\]', display: true},",
            "    {left: '\\\\(', right: '\\\\)', display: false}",
            "  ]});",
            "});",
            "</script>",
            "</body>",
            "</html>",
        ])

        output_path.write_text("\n".join(html_parts))
        return output_path

    def _export_pdf(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to PDF via LaTeX."""
        # First create LaTeX file
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_path = Path(tmpdir) / "document.tex"
            self._export_latex(sections, tex_path, **kwargs)

            # Try to compile with pdflatex
            try:
                result = subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", "-output-directory", tmpdir, str(tex_path)],
                    capture_output=True,
                    timeout=30,
                )

                # Run twice for references
                subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", "-output-directory", tmpdir, str(tex_path)],
                    capture_output=True,
                    timeout=30,
                )

                pdf_path = Path(tmpdir) / "document.pdf"
                if pdf_path.exists():
                    output_path.write_bytes(pdf_path.read_bytes())
                    return output_path
                else:
                    # Fall back to LaTeX output
                    output_path = output_path.with_suffix('.tex')
                    output_path.write_text(tex_path.read_text())
                    return output_path

            except (subprocess.TimeoutExpired, FileNotFoundError):
                # pdflatex not available, just save LaTeX
                output_path = output_path.with_suffix('.tex')
                output_path.write_text(tex_path.read_text())
                return output_path

    def _export_json(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to JSON."""
        data = {
            "title": kwargs.get('title', 'Document'),
            "author": kwargs.get('author', ''),
            "date": kwargs.get('date', datetime.now().isoformat()),
            "sections": sections,
            "metadata": {k: v for k, v in kwargs.items() if k not in ['title', 'author', 'date']},
        }
        output_path.write_text(json.dumps(data, indent=2, default=str))
        return output_path

    def _export_docx(self, sections: list[dict], output_path: Path, **kwargs) -> Path:
        """Export to Microsoft Word (.docx)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx required for .docx export: pip install python-docx")

        doc = Document()

        # Title
        if kwargs.get('title'):
            title_para = doc.add_heading(kwargs['title'], 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author and date
        if kwargs.get('author') or kwargs.get('date'):
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if kwargs.get('author'):
                meta.add_run(f"By {kwargs['author']}").italic = True
            if kwargs.get('date'):
                meta.add_run(f"\n{kwargs.get('date', datetime.now().strftime('%Y-%m-%d'))}")

        # Sections
        for section in sections:
            title = section.get('title', 'Section')
            content = section.get('content', '')
            level = section.get('level', 1)

            # Add heading
            doc.add_heading(title, level=min(level, 9))

            # Process content - handle markdown-like formatting
            self._add_docx_content(doc, content)

        doc.save(output_path)
        return output_path

    def _add_docx_content(self, doc, content: str):
        """Add content to docx, handling basic markdown."""
        import re
        from docx.shared import Pt

        lines = content.split('\n')
        current_para = None
        in_code_block = False
        code_lines = []

        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_lines)
                    para = doc.add_paragraph()
                    run = para.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    para.style = 'No Spacing'
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Headers
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            # List items
            elif line.strip().startswith('- '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line.strip()):
                text = re.sub(r'^\d+\. ', '', line.strip())
                doc.add_paragraph(text, style='List Number')
            # Regular paragraph
            elif line.strip():
                para = doc.add_paragraph()
                # Handle bold and italic
                self._add_formatted_text(para, line)
            # Empty line
            else:
                if current_para:
                    current_para = None

    def _add_formatted_text(self, para, text: str):
        """Add text with bold/italic formatting to paragraph."""
        import re

        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'

        for match in re.finditer(pattern, text):
            segment = match.group()
            if segment.startswith('**') and segment.endswith('**'):
                run = para.add_run(segment[2:-2])
                run.bold = True
            elif segment.startswith('*') and segment.endswith('*'):
                run = para.add_run(segment[1:-1])
                run.italic = True
            else:
                para.add_run(segment)

    def _markdown_to_latex(self, content: str) -> str:
        """Simple markdown to LaTeX conversion."""
        import re

        result = content

        # Escape special LaTeX characters (except those we'll handle)
        # Do this first before other transformations
        escape_chars = [('&', '\\&'), ('%', '\\%'), ('$', '\\$'), ('#', '\\#'), ('_', '\\_')]
        for char, escaped in escape_chars:
            # Don't escape if already escaped or inside code blocks
            result = re.sub(r'(?<!\\)' + re.escape(char), escaped, result)

        # Headers - use regex for proper matching
        result = re.sub(r'^### (.+)$', r'\\subsubsection*{\1}', result, flags=re.MULTILINE)
        result = re.sub(r'^## (.+)$', r'\\subsection*{\1}', result, flags=re.MULTILINE)

        # Bold (before italic to handle **text** vs *text*)
        result = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', result)

        # Italic
        result = re.sub(r'\*(.+?)\*', r'\\textit{\1}', result)

        # Code blocks
        result = re.sub(r'```\w*\n', r'\\begin{verbatim}\n', result)
        result = re.sub(r'```', r'\\end{verbatim}', result)

        # Inline code
        result = re.sub(r'`([^`]+)`', r'\\texttt{\1}', result)

        # Lists
        lines = result.split("\n")
        in_list = False
        new_lines = []

        for line in lines:
            if line.strip().startswith("- "):
                if not in_list:
                    new_lines.append("\\begin{itemize}")
                    in_list = True
                new_lines.append("\\item " + line.strip()[2:])
            else:
                if in_list:
                    new_lines.append("\\end{itemize}")
                    in_list = False
                new_lines.append(line)

        if in_list:
            new_lines.append("\\end{itemize}")

        return "\n".join(new_lines)

    def _markdown_to_html(self, content: str) -> str:
        """Simple markdown to HTML conversion."""
        result = content

        # Bold
        import re
        result = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', result)

        # Italic
        result = re.sub(r'\*(.+?)\*', r'<em>\1</em>', result)

        # Code blocks
        result = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code class="\1">\2</code></pre>', result, flags=re.DOTALL)

        # Inline code
        result = re.sub(r'`(.+?)`', r'<code>\1</code>', result)

        # Lists
        lines = result.split("\n")
        in_list = False
        new_lines = []

        for line in lines:
            if line.strip().startswith("- "):
                if not in_list:
                    new_lines.append("<ul>")
                    in_list = True
                new_lines.append(f"<li>{line.strip()[2:]}</li>")
            else:
                if in_list:
                    new_lines.append("</ul>")
                    in_list = False
                # Paragraphs
                if line.strip() and not line.startswith("<"):
                    new_lines.append(f"<p>{line}</p>")
                else:
                    new_lines.append(line)

        if in_list:
            new_lines.append("</ul>")

        return "\n".join(new_lines)

    def _get_html_styles(self) -> str:
        """Return CSS styles for HTML export."""
        return """
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 2rem;
            color: #333;
        }
        h1 { color: #2563eb; border-bottom: 2px solid #2563eb; padding-bottom: 0.5rem; }
        h2 { color: #1e40af; margin-top: 2rem; }
        h3 { color: #3730a3; }
        .author { color: #666; font-style: italic; }
        .date { color: #888; font-size: 0.9rem; }
        pre {
            background: #f4f4f5;
            padding: 1rem;
            border-radius: 4px;
            overflow-x: auto;
        }
        code {
            background: #f4f4f5;
            padding: 0.2rem 0.4rem;
            border-radius: 3px;
            font-family: 'Fira Code', 'Consolas', monospace;
        }
        pre code { background: none; padding: 0; }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 1rem 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 0.5rem;
            text-align: left;
        }
        th { background: #f4f4f5; }
        blockquote {
            border-left: 4px solid #2563eb;
            margin: 1rem 0;
            padding-left: 1rem;
            color: #555;
        }
        """


def quick_export(content: str, format: str = "markdown", output_path: str = None) -> str:
    """Quick helper to export content."""
    exporter = DocumentExporter()
    sections = [{"title": "Document", "content": content}]

    if output_path is None:
        output_path = f"output.{format}" if format != "pdf" else "output.pdf"

    result = exporter.export(sections, Path(output_path), format)
    return str(result)
