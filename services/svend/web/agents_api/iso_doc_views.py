"""ISO Document Creator API views.

Structured document authoring for ISO-compliant documents.
Documents can optionally be published to Document Control.
"""

import json
import re
import uuid
from io import BytesIO

from django.db import models
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from django.views.decorators.http import require_http_methods

from accounts.permissions import require_team

from .iso_document_types import ISO_DOCUMENT_TYPES
from .models import (
    ControlledDocument,
    DocumentRevision,
    ISODocument,
    ISOSection,
)

# =============================================================================
# Document Type Registry
# =============================================================================


@require_team
@require_http_methods(["GET"])
def list_types(request):
    """Return all available document types with their default sections."""
    types = {}
    for key, td in ISO_DOCUMENT_TYPES.items():
        types[key] = {
            "name": td["name"],
            "description": td["description"],
            "iso_clause": td.get("iso_clause", ""),
            "category": td.get("category", ""),
            "section_count": len(td["default_sections"]),
            "sections": [
                {"key": s["key"], "title": s["title"], "type": s["type"]}
                for s in td["default_sections"]
            ],
        }
    return JsonResponse({"types": types})


# =============================================================================
# Document CRUD
# =============================================================================


@require_team
@require_http_methods(["GET", "POST"])
def document_list_create(request):
    """List or create ISO documents."""
    user = request.user

    if request.method == "GET":
        docs = ISODocument.objects.filter(owner=user)

        # Filters
        status = request.GET.get("status")
        if status:
            docs = docs.filter(status=status)
        doc_type = request.GET.get("type")
        if doc_type:
            docs = docs.filter(document_type=doc_type)
        search = request.GET.get("search", "").strip()
        if search:
            docs = docs.filter(title__icontains=search)

        # Sort
        sort = request.GET.get("sort", "-updated_at")
        allowed_sorts = {
            "title",
            "-title",
            "updated_at",
            "-updated_at",
            "created_at",
            "-created_at",
            "status",
        }
        if sort in allowed_sorts:
            docs = docs.order_by(sort)

        docs = docs[:100]
        return JsonResponse({"documents": [d.to_dict() for d in docs]})

    # POST — create
    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    document_type = data.get("document_type", "")
    if document_type not in ISO_DOCUMENT_TYPES:
        return JsonResponse(
            {
                "error": f"Invalid document type: {document_type}",
                "valid_types": list(ISO_DOCUMENT_TYPES.keys()),
            },
            status=400,
        )

    type_def = ISO_DOCUMENT_TYPES[document_type]
    title = data.get("title", "").strip()
    if not title:
        title = f"New {type_def['name']}"

    doc = ISODocument.objects.create(
        owner=user,
        document_type=document_type,
        title=title,
        document_number=data.get("document_number", ""),
        iso_clause=type_def.get("iso_clause", ""),
        metadata=data.get("metadata", {}),
    )

    # Create default sections from the registry
    for i, sec_def in enumerate(type_def["default_sections"]):
        structured_data = sec_def.get("default_data", {})
        ISOSection.objects.create(
            document=doc,
            sort_order=i,
            section_type=sec_def["type"],
            section_key=sec_def["key"],
            title=sec_def["title"],
            structured_data=structured_data,
        )

    return JsonResponse(doc.to_dict_full(), status=201)


@require_team
@require_http_methods(["GET", "PUT", "DELETE"])
def document_detail(request, doc_id):
    """Get, update, or delete an ISO document."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)

    if request.method == "GET":
        return JsonResponse(doc.to_dict_full())

    if request.method == "DELETE":
        doc.delete()
        return JsonResponse({"ok": True})

    # PUT — update document-level fields
    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    updatable = ["title", "document_number", "status", "version", "iso_clause"]
    for field in updatable:
        if field in data:
            setattr(doc, field, data[field])

    if "metadata" in data:
        merged = doc.metadata or {}
        merged.update(data["metadata"])
        doc.metadata = merged

    doc.save()
    return JsonResponse(doc.to_dict_full())


# =============================================================================
# Section CRUD
# =============================================================================


@require_team
@require_http_methods(["POST"])
def section_create(request, doc_id):
    """Add a new section to a document."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    section_type = data.get("section_type", "paragraph")
    valid_types = [c[0] for c in ISOSection.SectionType.choices]
    if section_type not in valid_types:
        return JsonResponse(
            {"error": f"Invalid section_type: {section_type}"}, status=400
        )

    # Auto sort_order: append at end
    max_order = (
        doc.sections.order_by("-sort_order")
        .values_list("sort_order", flat=True)
        .first()
    )
    sort_order = (max_order or 0) + 1

    # Insert at specific position if requested
    if "after" in data:
        after_id = data["after"]
        try:
            after_sec = doc.sections.get(id=after_id)
            sort_order = after_sec.sort_order + 1
            # Shift subsequent sections
            doc.sections.filter(sort_order__gte=sort_order).update(
                sort_order=models.F("sort_order") + 1
            )
        except ISOSection.DoesNotExist:
            pass

    section = ISOSection.objects.create(
        document=doc,
        sort_order=sort_order,
        section_type=section_type,
        section_key=data.get("section_key", ""),
        title=data.get("title", ""),
        content=data.get("content", ""),
        structured_data=data.get("structured_data", {}),
        numbering=data.get("numbering", ""),
    )

    return JsonResponse(section.to_dict(), status=201)


@require_team
@require_http_methods(["PUT", "DELETE"])
def section_detail(request, doc_id, sec_id):
    """Update or delete a section."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)
    section = get_object_or_404(ISOSection, id=sec_id, document=doc)

    if request.method == "DELETE":
        section.delete()
        return JsonResponse({"ok": True})

    # PUT — partial update
    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    updatable = [
        "title",
        "content",
        "section_type",
        "numbering",
        "section_key",
        "image_caption",
    ]
    for field in updatable:
        if field in data:
            setattr(section, field, data[field])

    if "structured_data" in data:
        section.structured_data = data["structured_data"]

    if "embedded_media" in data:
        section.embedded_media = data["embedded_media"]

    # Image attachment by file_id
    if "image_id" in data:
        if data["image_id"]:
            from files.models import UserFile

            try:
                uf = UserFile.objects.get(id=data["image_id"], user=request.user)
                section.image = uf
            except UserFile.DoesNotExist:
                return JsonResponse({"error": "File not found"}, status=404)
        else:
            section.image = None

    section.save()
    return JsonResponse(section.to_dict())


@require_team
@require_http_methods(["POST"])
def section_reorder(request, doc_id):
    """Reorder sections within a document."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    order = data.get("order", [])
    if not order:
        return JsonResponse({"error": "order list required"}, status=400)

    for i, sec_id in enumerate(order):
        doc.sections.filter(id=sec_id).update(sort_order=i)

    return JsonResponse({"ok": True, "count": len(order)})


# =============================================================================
# Whiteboard Embed
# =============================================================================


@require_team
@require_http_methods(["POST"])
def embed_whiteboard(request, doc_id, sec_id):
    """Embed a whiteboard export (SVG or PNG) into a section."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)
    section = get_object_or_404(ISOSection, id=sec_id, document=doc)

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    room_code = data.get("room_code", "").strip().upper()
    if not room_code:
        return JsonResponse({"error": "room_code required"}, status=400)

    fmt = data.get("format", "svg")
    if fmt not in ("svg", "png"):
        return JsonResponse({"error": "format must be 'svg' or 'png'"}, status=400)

    from .models import Board

    try:
        board = Board.objects.get(room_code=room_code)
    except Board.DoesNotExist:
        return JsonResponse({"error": "Whiteboard not found"}, status=404)

    from .whiteboard_views import _generate_svg

    svg_content, width, height = _generate_svg(board, theme="light")
    if not svg_content:
        return JsonResponse({"error": "Whiteboard is empty"}, status=400)

    media_entry = {
        "id": str(uuid.uuid4())[:8],
        "board_name": board.name,
        "room_code": board.room_code,
        "format": fmt,
        "width": width,
        "height": height,
    }

    if fmt == "svg":
        media_entry["data"] = svg_content
    else:
        # Convert SVG to PNG via cairosvg
        import cairosvg

        png_bytes = cairosvg.svg2png(
            bytestring=svg_content.encode("utf-8"),
            output_width=min(width, 1200),
        )
        import base64

        media_entry["data"] = base64.b64encode(png_bytes).decode("utf-8")

    media = section.embedded_media or []
    media.append(media_entry)
    section.embedded_media = media
    section.save(update_fields=["embedded_media"])

    return JsonResponse(
        {
            "ok": True,
            "media_id": media_entry["id"],
            "format": fmt,
            "board_name": board.name,
        }
    )


# =============================================================================
# Export — PDF
# =============================================================================


@require_team
@require_http_methods(["GET"])
def export_pdf(request, doc_id):
    """Export document as PDF via weasyprint."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)
    sections = list(doc.sections.order_by("sort_order"))
    type_def = ISO_DOCUMENT_TYPES.get(doc.document_type, {})

    html_string = render_to_string(
        "iso_document_print.html",
        {
            "document": doc,
            "sections": sections,
            "type_def": type_def,
        },
    )

    pdf_buffer = BytesIO()
    try:
        from weasyprint import HTML

        HTML(string=html_string, base_url="https://svend.ai").write_pdf(pdf_buffer)
    except Exception:
        import os
        import subprocess
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w") as f:
            f.write(html_string)
            html_path = f.name
        pdf_path = html_path.replace(".html", ".pdf")
        try:
            subprocess.run(
                ["wkhtmltopdf", "--quiet", html_path, pdf_path],
                check=True,
                timeout=30,
            )
            with open(pdf_path, "rb") as pf:
                pdf_buffer.write(pf.read())
        finally:
            for p in (html_path, pdf_path):
                try:
                    os.unlink(p)
                except OSError:
                    pass

    pdf_buffer.seek(0)
    safe_name = re.sub(r"[^\w\-.]", "_", doc.title)[:60] or "document"
    response = HttpResponse(pdf_buffer.read(), content_type="application/pdf")
    response["Content-Disposition"] = f'inline; filename="{safe_name}.pdf"'
    return response


# =============================================================================
# Export — Word (.docx)
# =============================================================================


@require_team
@require_http_methods(["GET"])
def export_docx(request, doc_id):
    """Export document as Word .docx via python-docx."""
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)
    sections = list(doc.sections.order_by("sort_order"))
    type_def = ISO_DOCUMENT_TYPES.get(doc.document_type, {})

    from docx import Document as DocxDocument

    docx = DocxDocument()

    # Title page
    docx.add_heading(doc.title, level=0)
    if doc.document_number:
        docx.add_paragraph(f"Document Number: {doc.document_number}")
    docx.add_paragraph(f"Type: {type_def.get('name', doc.document_type)}")
    docx.add_paragraph(f"Version: {doc.version}")
    if doc.iso_clause:
        docx.add_paragraph(f"ISO Clause: {doc.iso_clause}")
    meta = doc.metadata or {}
    if meta.get("prepared_by"):
        docx.add_paragraph(f"Prepared By: {meta['prepared_by']}")
    if meta.get("effective_date"):
        docx.add_paragraph(f"Effective Date: {meta['effective_date']}")
    docx.add_page_break()

    # Render sections
    for section in sections:
        _render_section_docx(docx, section)

    buffer = BytesIO()
    docx.save(buffer)
    buffer.seek(0)

    safe_name = re.sub(r"[^\w\-.]", "_", doc.title)[:60] or "document"
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.docx"'
    return response


def _render_section_docx(docx, section):
    """Render a single ISOSection into a python-docx Document."""
    from docx.shared import Inches

    st = section.section_type
    sd = section.structured_data or {}

    if st == "heading":
        level = sd.get("level", 1)
        if section.numbering:
            docx.add_heading(
                f"{section.numbering} {section.title}", level=min(level, 3)
            )
        else:
            docx.add_heading(section.title, level=min(level, 3))
        if section.content:
            docx.add_paragraph(section.content)

    elif st == "paragraph":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        if section.content:
            docx.add_paragraph(section.content)

    elif st == "definition":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        items = sd.get("items", [])
        for item in items:
            para = docx.add_paragraph()
            run = para.add_run(f"{item.get('term', '')}: ")
            run.bold = True
            para.add_run(item.get("definition", ""))

    elif st == "reference":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        refs = sd.get("refs", [])
        for ref in refs:
            parts = []
            if ref.get("document_number"):
                parts.append(ref["document_number"])
            if ref.get("title"):
                parts.append(ref["title"])
            if ref.get("clause"):
                parts.append(f"(clause {ref['clause']})")
            docx.add_paragraph(" — ".join(parts), style="List Bullet")
        if section.content:
            docx.add_paragraph(section.content)

    elif st == "table":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        columns = sd.get("columns", [])
        rows = sd.get("rows", [])
        if columns:
            table = docx.add_table(rows=1 + len(rows), cols=len(columns))
            table.style = "Table Grid"
            # Header row
            for j, col in enumerate(columns):
                cell = table.rows[0].cells[j]
                cell.text = str(col)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            # Data rows
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    if j < len(columns):
                        table.rows[i + 1].cells[j].text = str(val)
            docx.add_paragraph()  # spacing

    elif st == "checklist":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        items = sd.get("items", [])
        for item in items:
            check = "\u2611" if item.get("checked") else "\u2610"
            docx.add_paragraph(f"{check} {item.get('text', '')}", style="List Bullet")

    elif st == "signature_block":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        signers = sd.get("signers", [])
        if signers:
            table = docx.add_table(rows=len(signers) + 1, cols=4)
            table.style = "Table Grid"
            for j, hdr in enumerate(["Role", "Name", "Signature", "Date"]):
                cell = table.rows[0].cells[j]
                cell.text = hdr
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for i, signer in enumerate(signers):
                table.rows[i + 1].cells[0].text = signer.get("role", "")
                table.rows[i + 1].cells[1].text = signer.get("name", "")
                table.rows[i + 1].cells[2].text = ""  # blank for signature
                table.rows[i + 1].cells[3].text = signer.get("date", "")
            docx.add_paragraph()

    elif st == "image":
        if section.title:
            prefix = f"{section.numbering} " if section.numbering else ""
            docx.add_heading(f"{prefix}{section.title}", level=2)
        if section.image_id:
            try:
                docx.add_picture(section.image.file.path, width=Inches(5))
            except Exception:
                docx.add_paragraph("[Image could not be embedded]")
        if section.image_caption:
            p = docx.add_paragraph(section.image_caption)
            p.style = "Caption" if "Caption" in [s.name for s in docx.styles] else None

    # Embedded media (whiteboard exports)
    media = section.embedded_media or []
    for m in media:
        if m.get("format") == "png" and m.get("data"):
            import base64

            img_bytes = base64.b64decode(m["data"])
            img_buffer = BytesIO(img_bytes)
            try:
                docx.add_picture(img_buffer, width=Inches(5))
                if m.get("board_name"):
                    docx.add_paragraph(f"Source: {m['board_name']}")
            except Exception:
                docx.add_paragraph(f"[Whiteboard: {m.get('board_name', 'unknown')}]")


# =============================================================================
# Publish to Document Control
# =============================================================================


@require_team
@require_http_methods(["POST"])
def publish_to_doc_control(request, doc_id):
    """Publish an ISODocument to the Document Control register.

    Creates a ControlledDocument from the authored content, or updates
    the existing linked one if re-publishing.
    """
    doc = get_object_or_404(ISODocument, id=doc_id, owner=request.user)

    if doc.status != ISODocument.Status.FINAL:
        return JsonResponse(
            {
                "error": "Document must be in 'Final' status to publish to Document Control",
            },
            status=400,
        )

    # Render sections to content
    sections = doc.sections.order_by("sort_order")
    content_parts = []
    for s in sections:
        if s.title:
            prefix = f"{s.numbering} " if s.numbering else ""
            content_parts.append(f"## {prefix}{s.title}")
        if s.content:
            content_parts.append(s.content)
        if s.section_type == "table":
            cols = (s.structured_data or {}).get("columns", [])
            rows = (s.structured_data or {}).get("rows", [])
            if cols:
                content_parts.append("| " + " | ".join(str(c) for c in cols) + " |")
                content_parts.append("| " + " | ".join("---" for _ in cols) + " |")
                for row in rows:
                    content_parts.append("| " + " | ".join(str(v) for v in row) + " |")
        if s.section_type == "checklist":
            for item in (s.structured_data or {}).get("items", []):
                check = "[x]" if item.get("checked") else "[ ]"
                content_parts.append(f"- {check} {item.get('text', '')}")
        if s.section_type == "definition":
            for item in (s.structured_data or {}).get("items", []):
                content_parts.append(
                    f"**{item.get('term', '')}**: {item.get('definition', '')}"
                )
        if s.section_type == "signature_block":
            for signer in (s.structured_data or {}).get("signers", []):
                content_parts.append(
                    f"_{signer.get('role', '')}_: ____________________  Date: ________"
                )

    content = "\n\n".join(content_parts)
    type_def = ISO_DOCUMENT_TYPES.get(doc.document_type, {})

    if doc.controlled_document_id:
        # Re-publish: snapshot current, update content
        cd = doc.controlled_document
        DocumentRevision.objects.create(
            document=cd,
            version=cd.current_version,
            content_snapshot=cd.content,
            change_summary="Re-published from ISO Document Creator",
            changed_by=request.user,
        )
        cd.content = content
        cd.title = doc.title
        cd.document_number = doc.document_number
        cd.iso_clause = doc.iso_clause
        # Bump version
        try:
            parts = cd.current_version.split(".")
            parts[-1] = str(int(parts[-1]) + 1)
            cd.current_version = ".".join(parts)
        except (ValueError, IndexError):
            cd.current_version = cd.current_version + ".1"
        cd.save()
    else:
        # First publish
        cd = ControlledDocument.objects.create(
            owner=request.user,
            title=doc.title,
            document_number=doc.document_number,
            category=type_def.get("category", ""),
            iso_clause=doc.iso_clause,
            content=content,
            current_version=doc.version,
        )
        doc.controlled_document = cd
        doc.save(update_fields=["controlled_document"])

    return JsonResponse(
        {
            "ok": True,
            "controlled_document": cd.to_dict(),
        }
    )
